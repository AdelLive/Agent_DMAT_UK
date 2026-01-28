import streamlit as st
import os
from pathlib import Path
import hashlib

from PyPDF2 import PdfReader
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from openai import OpenAI

# Import Word avec gestion d'erreur
try:
    from docx import Document
    HAS_WORD = True
except ImportError:
    HAS_WORD = False

st.set_page_config(page_title="Agent Docs", page_icon="🤖")

DATA_DIR = Path("data")
DATA_DIR.mkdir(exist_ok=True)

class SimpleRAG:
    def __init__(self):
        self.vectorizer = TfidfVectorizer()
        self.documents = []
        self.vectors = None
        self.load_data()
    
    def load_data(self):
        if Path("data/metadata.csv").exists():
            df = pd.read_csv("data/metadata.csv")
            self.documents = df.to_dict('records')
            if Path("data/vectors.npz").exists():
                data = np.load("data/vectors.npz", allow_pickle=True)
                self.vectors = data['vectors']
                if self.documents:
                    texts = [d['content'] for d in self.documents]
                    self.vectorizer.fit(texts)
    
    def save_data(self):
        if self.documents:
            df = pd.DataFrame(self.documents)
            df.to_csv("data/metadata.csv", index=False)
            if self.vectors is not None:
                np.savez("data/vectors.npz", vectors=self.vectors)
    
    def extract_text(self, file_path, file_type):
        text = ""
        try:
            if file_type == "pdf":
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            
            elif file_type == "docx" and HAS_WORD:
                doc = Document(file_path)
                # Récupère tout le texte des paragraphes
                full_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)
                # Récupère aussi le texte des tableaux
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text)
                        if row_text:
                            full_text.append(" | ".join(row_text))
                text = "\n".join(full_text)
            
            elif file_type in ["txt", "md"]:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif file_type == "csv":
                df = pd.read_csv(file_path)
                text = df.to_string()
            elif file_type == "xlsx":
                df = pd.read_excel(file_path)
                text = df.to_string()
            return text
        except Exception as e:
            st.error(f"Erreur lecture {file_type}: {e}")
            return ""
    
    def add_document(self, file_obj, file_name):
        ext = file_name.split('.')[-1].lower()
        
        with open(DATA_DIR / file_name, "wb") as f:
            f.write(file_obj.getvalue())
        
        content = self.extract_text(DATA_DIR / file_name, ext)
        
        if not content or len(content.strip()) == 0:
            st.warning(f"⚠️ Fichier vide ou non lisible")
            return 0
        
        chunks = [content[i:i+1000] for i in range(0, len(content), 1000)]
        
        for i, chunk in enumerate(chunks):
            self.documents.append({
                'file_name': file_name,
                'content': chunk,
                'chunk_index': i
            })
        
        texts = [d['content'] for d in self.documents]
        self.vectors = self.vectorizer.fit_transform(texts)
        self.save_data()
        return len(chunks)
    
    def search(self, query, k=3):
        if not self.documents or self.vectors is None:
            return []
        query_vec = self.vectorizer.transform([query])
        similarities = cosine_similarity(query_vec, self.vectors).flatten()
        top_indices = similarities.argsort()[-k:][::-1]
        return [self.documents[i] for i in top_indices if similarities[i] > 0.1]

if 'rag' not in st.session_state:
    st.session_state.rag = SimpleRAG()
if 'messages' not in st.session_state:
    st.session_state.messages = []

st.title("🤖 Agent Documents")

with st.sidebar:
    st.header("Configuration")
    
    api_key = st.text_input("Clé API Kimi", type="password")
    
    if not api_key:
        st.warning("Entre ta clé API")
        st.stop()
    
    client = OpenAI(api_key=api_key, base_url="https://api.moonshot.cn/v1")
    st.success("Prêt")
    
    st.divider()
    
    # Types acceptés selon si python-docx est installé
    file_types = ["pdf", "txt", "csv", "xlsx"]
    if HAS_WORD:
        file_types.append("docx")
        st.caption("✅ Support Word activé")
    else:
        st.caption("❌ pip3 install python-docx pour Word")
    
    uploaded_file = st.file_uploader("Fichier", type=file_types)
    
    if uploaded_file:
        with st.spinner("Traitement..."):
            n = st.session_state.rag.add_document(uploaded_file, uploaded_file.name)
            if n > 0:
                st.success(f"{n} sections ajoutées")
                st.rerun()
    
    if st.session_state.rag.documents:
        files = list(set([d['file_name'] for d in st.session_state.rag.documents]))
        st.write(f"**{len(files)} fichier(s)**")

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])

if prompt := st.chat_input("Question?"):
    if not st.session_state.rag.documents:
        st.error("Upload un fichier d'abord")
    else:
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.write(prompt)
        
        docs = st.session_state.rag.search(prompt)
        
        if not docs:
            response = "Aucune info trouvée."
        else:
            context = "\n".join([d['content'][:800] for d in docs])
            
            try:
                resp = client.chat.completions.create(
                    model="moonshot-v1-8k",
                    messages=[
                        {"role": "system", "content": f"Réponds basé sur:\n{context}"},
                        {"role": "user", "content": prompt}
                    ]
                )
                response = resp.choices[0].message.content
            except Exception as e:
                response = f"Erreur: {e}"
        
        st.session_state.messages.append({"role": "assistant", "content": response})
        with st.chat_message("assistant"):
            st.write(response)